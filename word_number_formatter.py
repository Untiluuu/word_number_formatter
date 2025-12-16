"""
Word文档数字格式化工具
将Word文档中的所有数字格式化为千分位，保持原有小数位数不变
例如：1000 -> 1,000, 10000.123 -> 10,000.123
"""

import re
import os
import sys
import webbrowser
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading


class WordNumberFormatter:
    """Word文档数字格式化处理器"""
    
    def __init__(self):
        # 匹配数字（包括负数），但排除日期格式和百分比
        # 负数格式：-10000.00
        # 不使用\b因为它不支持负号，改用(?<!\w)和(?!\w)
        self.number_pattern = re.compile(
            r'(?<![/\w])-?\d+(?:\.\d+)?(?!\w|[-/年月日时分秒%])'
        )
    
    def format_number(self, num_str):
        """
        格式化数字为千分位，保留原有小数位数（支持负数）
        :param num_str: 数字字符串
        :return: 格式化后的数字字符串
        """
        try:
            # 检查是否是负数
            is_negative = num_str.startswith('-')
            if is_negative:
                num_str = num_str[1:]  # 移除负号
            
            # 检查是否包含小数点
            if '.' in num_str:
                # 分离整数和小数部分
                integer_part, decimal_part = num_str.split('.')
                # 只对整数部分添加千分位
                integer_num = int(integer_part)
                formatted_integer = "{:,}".format(integer_num)
                # 保留原有的小数部分，不做处理
                result = f"{formatted_integer}.{decimal_part}"
            else:
                # 没有小数点，只处理整数，不添加小数位
                num = int(num_str)
                result = "{:,}".format(num)
            
            # 如果是负数，添加负号
            if is_negative:
                result = '-' + result
            
            return result
        except ValueError:
            return num_str if not is_negative else '-' + num_str
    
    def process_text(self, text):
        """
        处理文本中的所有数字（排除日期格式）
        :param text: 原始文本
        :return: 处理后的文本
        """
        def replace_number(match):
            num_str = match.group()
            # 再次检查：如果数字是4位数且在1900-2999之间，可能是年份，需要更谨慎处理
            # 检查前后文是否有日期相关字符
            start, end = match.span()
            before = text[max(0, start-1):start] if start > 0 else ''
            after = text[end:end+1] if end < len(text) else ''
            
            # 如果前后有日期相关字符或百分号，不处理
            date_chars = ['年', '月', '日', '时', '分', '秒', '-', '/', '.', '%']
            if before in date_chars or after in date_chars:
                return num_str
            
            # 如果是4位数字且可能是年份（1900-2999），检查后续是否跟年月日
            if len(num_str) == 4 and num_str.isdigit():
                year_num = int(num_str)
                if 1900 <= year_num <= 2999:
                    # 检查后面是否紧跟年月日或日期分隔符
                    next_chars = text[end:end+2] if end+2 <= len(text) else text[end:]
                    if any(c in next_chars for c in ['年', '-', '/']):
                        return num_str
            
            return self.format_number(num_str)
        
        return self.number_pattern.sub(replace_number, text)
    
    def process_paragraph(self, paragraph):
        """
        处理段落中的数字，保持原有格式，并添加黄色底色
        :param paragraph: 段落对象
        """
        # 先获取完整的段落文本
        full_text = paragraph.text
        
        # 如果段落为空或没有run，直接返回
        if not full_text or not paragraph.runs:
            return
        
        # 处理完整文本，获取格式化后的结果
        formatted_full_text = self.process_text(full_text)
        
        # 如果文本没有变化，不需要处理
        if full_text == formatted_full_text:
            return
        
        # 文本发生了变化，需要更新所有runs
        # 最简单的方式：清空所有runs，然后在第一个run中放入格式化后的文本
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                # 第一个run放入完整的格式化文本
                run.text = formatted_full_text
                # 添加黄色底色标记
                shading_elm = run._element.get_or_add_rPr()
                shading = shading_elm.find(qn('w:shd'))
                if shading is None:
                    shading = shading_elm.makeelement(qn('w:shd'))
                    shading_elm.append(shading)
                shading.set(qn('w:fill'), 'FFFF00')  # 黄色底色
            else:
                # 其他runs清空
                run.text = ''
    
    def process_table_cell(self, cell):
        """
        处理表格单元格中的数字
        :param cell: 单元格对象
        """
        for paragraph in cell.paragraphs:
            self.process_paragraph(paragraph)
    
    def process_document(self, input_file, output_file=None):
        """
        处理Word文档
        :param input_file: 输入文件路径
        :param output_file: 输出文件路径，如果为None则覆盖原文件
        :return: 处理结果信息
        """
        try:
            # 打开文档
            doc = Document(input_file)
            
            paragraph_count = 0
            table_count = 0
            
            # 处理所有段落
            for paragraph in doc.paragraphs:
                self.process_paragraph(paragraph)
                paragraph_count += 1
            
            # 处理所有表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.process_table_cell(cell)
                table_count += 1
            
            # 保存文档
            if output_file is None:
                output_file = input_file
            
            doc.save(output_file)
            
            return {
                'success': True,
                'message': f'处理完成！\n处理段落数：{paragraph_count}\n处理表格数：{table_count}',
                'output_file': output_file
            }
        
        except Exception as e:
            return {
                'success': False,
                'message': f'处理失败：{str(e)}',
                'output_file': None
            }


class WordFormatterGUI:
    """Word文档数字格式化工具图形界面"""
    
    VERSION = "V1.0.1"
    
    def __init__(self):
        self.window = tk.Tk()
        self.window.title("Word文档数字格式化工具")
        self.window.geometry("600x520")
        self.window.resizable(False, False)
        
        self.formatter = WordNumberFormatter()
        self.input_file = None
        self.output_file = None
        
        self.setup_ui()
    
    def setup_ui(self):
        """设置用户界面"""
        # 标题
        title_frame = tk.Frame(self.window, bg="#4A90E2", height=60)
        title_frame.pack(fill=tk.X)
        
        title_label = tk.Label(
            title_frame,
            text=f"Word文档数字格式化工具 {self.VERSION}",
            font=("微软雅黑", 16, "bold"),
            bg="#4A90E2",
            fg="white"
        )
        title_label.pack(pady=15)
        
        # 主容器
        main_frame = tk.Frame(self.window, padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 说明文字
        desc_label = tk.Label(
            main_frame,
            text="将Word文档中的所有数字格式化为千分位\n例如：1000 → 1,000，10000.123 → 10,000.123\n✨ 修改的数字会添加黄色底色标记，便于查找",
            font=("微软雅黑", 10),
            fg="#666666",
            justify=tk.LEFT
        )
        desc_label.pack(pady=(0, 20))
        
        # 输入文件选择
        input_frame = tk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=10)
        
        tk.Label(input_frame, text="输入文件：", font=("微软雅黑", 10)).pack(side=tk.LEFT)
        
        self.input_entry = tk.Entry(input_frame, font=("微软雅黑", 9), state='readonly')
        self.input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10)
        
        tk.Button(
            input_frame,
            text="选择文件",
            command=self.select_input_file,
            font=("微软雅黑", 9),
            bg="#4A90E2",
            fg="white",
            cursor="hand2"
        ).pack(side=tk.RIGHT)
        
        # 输出选项
        output_frame = tk.Frame(main_frame)
        output_frame.pack(fill=tk.X, pady=10)
        
        self.overwrite_var = tk.BooleanVar(value=True)
        tk.Radiobutton(
            output_frame,
            text="覆盖原文件",
            variable=self.overwrite_var,
            value=True,
            font=("微软雅黑", 10),
            command=self.toggle_output_option
        ).pack(anchor=tk.W)
        
        tk.Radiobutton(
            output_frame,
            text="另存为新文件",
            variable=self.overwrite_var,
            value=False,
            font=("微软雅黑", 10),
            command=self.toggle_output_option
        ).pack(anchor=tk.W, pady=5)
        
        # 输出文件选择（默认隐藏）
        self.output_select_frame = tk.Frame(main_frame)
        self.output_select_frame.pack(fill=tk.X, pady=10)
        
        tk.Label(self.output_select_frame, text="输出文件：", font=("微软雅黑", 10)).pack(side=tk.LEFT)
        
        self.output_entry = tk.Entry(self.output_select_frame, font=("微软雅黑", 9), state='readonly')
        self.output_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10)
        
        tk.Button(
            self.output_select_frame,
            text="选择位置",
            command=self.select_output_file,
            font=("微软雅黑", 9),
            bg="#4A90E2",
            fg="white",
            cursor="hand2"
        ).pack(side=tk.RIGHT)
        
        self.output_select_frame.pack_forget()  # 初始隐藏
        
        # 进度条
        self.progress_frame = tk.Frame(main_frame)
        self.progress_frame.pack(fill=tk.X, pady=20)
        
        self.progress = ttk.Progressbar(
            self.progress_frame,
            mode='indeterminate',
            length=500
        )
        self.progress.pack()
        self.progress_frame.pack_forget()  # 初始隐藏
        
        # 处理按钮
        self.process_btn = tk.Button(
            main_frame,
            text="开始处理",
            command=self.process_file,
            font=("微软雅黑", 12, "bold"),
            bg="#5CB85C",
            fg="white",
            width=20,
            height=2,
            cursor="hand2"
        )
        self.process_btn.pack(pady=20)
        
        # 作者信息栏
        author_frame = tk.Frame(self.window, bg="#F5F5F5", height=30)
        author_frame.pack(side=tk.BOTTOM, fill=tk.X)
        
        author_label = tk.Label(
            author_frame,
            text="作者：taoPeng",
            font=("微软雅黑", 9),
            bg="#F5F5F5",
            fg="#666666"
        )
        author_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        github_label = tk.Label(
            author_frame,
            text="GitHub: https://github.com/Untiluuu",
            font=("微软雅黑", 9),
            bg="#F5F5F5",
            fg="#4A90E2",
            cursor="hand2"
        )
        github_label.pack(side=tk.LEFT, padx=5, pady=5)
        github_label.bind("<Button-1>", lambda e: self.open_github())
        
        # 状态栏
        self.status_label = tk.Label(
            self.window,
            text="就绪",
            font=("微软雅黑", 9),
            bg="#E8E8E8",
            anchor=tk.W,
            padx=10
        )
        self.status_label.pack(side=tk.BOTTOM, fill=tk.X)
    
    def toggle_output_option(self):
        """切换输出选项"""
        if self.overwrite_var.get():
            self.output_select_frame.pack_forget()
        else:
            self.output_select_frame.pack(fill=tk.X, pady=10)
    
    def select_input_file(self):
        """选择输入文件"""
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.input_file = filename
            self.input_entry.config(state='normal')
            self.input_entry.delete(0, tk.END)
            self.input_entry.insert(0, filename)
            self.input_entry.config(state='readonly')
            self.status_label.config(text=f"已选择文件：{os.path.basename(filename)}")
    
    def select_output_file(self):
        """选择输出文件"""
        if not self.input_file:
            messagebox.showwarning("警告", "请先选择输入文件！")
            return
        
        input_dir = os.path.dirname(self.input_file)
        input_name = os.path.basename(self.input_file)
        name_without_ext = os.path.splitext(input_name)[0]
        default_name = f"{name_without_ext}_格式化.docx"
        
        filename = filedialog.asksaveasfilename(
            title="保存为",
            initialdir=input_dir,
            initialfile=default_name,
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.output_file = filename
            self.output_entry.config(state='normal')
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, filename)
            self.output_entry.config(state='readonly')
    
    def process_file(self):
        """处理文件"""
        if not self.input_file:
            messagebox.showwarning("警告", "请先选择输入文件！")
            return
        
        # 如果选择了另存为新文件，但还没有选择输出位置，则自动弹出选择对话框
        if not self.overwrite_var.get() and not self.output_file:
            self.select_output_file()
            # 如果用户取消了选择，则不继续处理
            if not self.output_file:
                return
        
        # 禁用按钮，显示进度条
        self.process_btn.config(state='disabled')
        self.progress_frame.pack(fill=tk.X, pady=20)
        self.progress.start(10)
        self.status_label.config(text="正在处理...")
        
        # 在新线程中处理文件
        thread = threading.Thread(target=self._process_file_thread)
        thread.daemon = True
        thread.start()
    
    def _process_file_thread(self):
        """处理文件的线程函数"""
        output = self.output_file if not self.overwrite_var.get() else None
        result = self.formatter.process_document(self.input_file, output)
        
        # 在主线程中更新UI
        self.window.after(0, self._process_complete, result)
    
    def _process_complete(self, result):
        """处理完成后的回调"""
        self.progress.stop()
        self.progress_frame.pack_forget()
        self.process_btn.config(state='normal')
        
        if result['success']:
            messagebox.showinfo("成功", result['message'])
            self.status_label.config(text="处理完成")
        else:
            messagebox.showerror("错误", result['message'])
            self.status_label.config(text="处理失败")
    
    def open_github(self):
        """打开GitHub主页"""
        webbrowser.open("https://github.com/Untiluuu")
    
    def run(self):
        """运行应用程序"""
        self.window.mainloop()


def main():
    """主函数"""
    app = WordFormatterGUI()
    app.run()


if __name__ == "__main__":
    main()
